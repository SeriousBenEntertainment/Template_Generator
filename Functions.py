### `Functions.py`
### ❄️ Template Generator
### Open-Source, hosted on https://github.com/SeriousBenEntertainment/Template_Generator
### Please reach out to ben@seriousbenentertainment.org for any questions
### Loading needed Python libraries
import streamlit as st
import pandas as pd
import datetime
import requests
import os
import warnings
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
warnings.filterwarnings(
    action='ignore',
    category=UserWarning,
    module='Minio'
)
from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from minio import Minio
from minio.error import S3Error
from typing import Any, List, Mapping, Optional
from langchain_core.callbacks.manager import CallbackManagerForLLMRun
from langchain_core.language_models.llms import LLM
from snowflake.snowpark import Session
from snowflake.cortex import Complete

# Cortex
class Cortex(LLM):
    session: Session = None

    model: str = "mistral-large"

    @property
    def _llm_type(self) -> str:
        return "cortex"

    def _call(
        self,
        prompt: str,
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> str:
        if stop is not None:
            raise ValueError("stop kwargs are not permitted.")

        response = Complete(
            model=self.model,
            prompt=prompt,
            session=self.session,
            stream=False
        )
        return response.replace("AI:", "")

    @property
    def _identifying_params(self) -> Mapping[str, Any]:
        """Get the identifying parameters."""
        return {
            "session": self.session, 
            "model": self.model
        }

# Establish MiniO session
def connect_to_minio(endpoint_url, access_key, secret_key, secure):
    try:
        client = Minio(endpoint_url,
        access_key=access_key,
        secret_key=secret_key,
        secure=secure, # Using HTTP or HTTPS
        cert_check=False) 
        return client
    except S3Error as e:
        st.error(f"Error: {e}")
    return None

# Function to list buckets
def list_buckets(minio_client):
    try:
        buckets = minio_client.list_buckets()
        return [
                    bucket.name.replace('-', ' ').title()
                    for bucket in buckets
                ]
    except S3Error as e:
        st.error(f"Error: {e}")
    return 

# Function to list objects in a bucket
def list_objects(minio_client, bucket_name):
    bucket_name = bucket_name.lower().replace(' ', '-')
    try:
        objects = minio_client.list_objects(bucket_name, recursive=True)
        return [
                    obj.object_name 
                    for obj in objects
                ]
    except S3Error as e:
        st.error(f"Error: {e}")
    return

def upload_files(minio_client, bucket_name, files):
    bucket_name = bucket_name.lower().replace(' ', '-')
    for file in files:
        # Read file
        file_content = file.read()
        
        # Upload to MinIO
        minio_client.put_object(
            bucket_name,
            file.name,
            BytesIO(file_content),
            len(file_content)
        )

# Establish Snowflake session
@st.cache_resource
def create_session():
    session = Session.builder.configs(st.secrets.snowflake).create()
    try:
        session.use_role(st.secrets.snowflake["role"])
        session.sql(f'USE WAREHOUSE "{st.secrets.snowflake["warehouse"]}"')
        session.use_database(st.secrets.snowflake["database"])
        session.use_schema(st.secrets.snowflake["schema"])
    except Exception as e:
        st.error(f"Error: {e}")
    return session

# Write data table
def write_data(session, data, table_name, database, schema):
    # Write data to table
    session.write_pandas(data, table_name=table_name, database=database, schema=schema, overwrite=True)
    st.success("Daten erfolgreich geschrieben.")

def uploading_files(session, stage_name, files):
    for file in files:
        # Read file(s)
        file_content = file.read()
        
        # Upload to Snowflake
        session.file.put_stream(
            BytesIO(file_content),
            f"{stage_name.replace(' ', '_')}/{file.name}",
            auto_compress=False,
            overwrite=True,
        )
        
# Load data table
def load_data(_session, table_name):
    # Read in data table
    table = _session.table(table_name)

    # Collect the results. This will run the query and download the data
    table = table.collect()
    st.success("Daten erfolgreich gelesen.")
    return pd.DataFrame(table)

# List stages
def list_stages(session):
    try:
        stages = session.sql("SHOW STAGES;").collect()
        return [
                    stage.name.replace('_', ' ').title()
                    for stage in stages
                ]
    except S3Error as e:
        st.error(f"Error: {e}")
    return

# Function to list objects in a Stage
def list_files(session, stage_name):
    try:
        query = f"LIST @{stage_name.replace(' ', '_')}/;"
        files = session.sql(query).collect()
        return [
                    os.path.basename(file.name)
                    for file in files
                ]
    except S3Error as e:
        st.error(f"Error: {e}")
    return

# Web Scraper
def web_scraper(url):
    info_page = requests.get(url)
    info_soup = BeautifulSoup(info_page.content, 'html.parser')
    info = info_soup.get_text()
    info = info.replace('\n', ' ')
    return info

# Word export
def export_doc(data, cloud, last_chapter, paragraph_of_summary, table_of_glossar, table_of_stakeholders, table_of_attachments, table_of_contents):
    document = Document()

    # Adding Image
    document.add_picture('images/gwq_logo_header.png')
    
    # Adding centered text
    centered_text = f"""\n\n\n
                        {cloud} 
                        \n\n
                        Anforderungen 
                        Regulatorik
                        """
    paragraph = document.add_paragraph(centered_text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(24)
    centered_date = f"""\n\n
                        {datetime.date.today()}"""
    paragraph = document.add_paragraph(centered_date)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(14)
    document.add_page_break()
    
    # Adding table of changes
    def set_table_borders(table):
        tbl = table._element
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Border size
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Border color
            tblBorders.append(border)
        tbl.tblPr.append(tblBorders)
    paragraph = document.add_paragraph('Änderungshistorie des Dokumentes')
    table = document.add_table(rows=5, cols=7)
    set_table_borders(table)
    hdr_cells = table.rows[0].cells
    headers = ['Nr.', 'Datum', 'Version', 'Kapitel', 'Beschreibung der Änderung', 'Autor', 'Bearbeitungszustand']
    for i, header in enumerate(headers):
        run = hdr_cells[i].add_paragraph().add_run(header)
        run.bold = True
        run.font.size = Pt(10)
    row_cells = table.rows[1].cells
    contents = [str(1), str(datetime.date.today()), '1.0', 'alle', 'Erstellung des Dokuments', 'Groß, Benjamin', 'Erl.']
    for i, content in enumerate(contents):
        run = row_cells[i].add_paragraph().add_run(content)
        run.font.size = Pt(10)
    document.add_page_break()

    # Adding Table of Contents
    if table_of_contents:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run('Inhaltsverzeichnis')
        run.font.size = Pt(16)

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:t')
        fldChar3.text = " (Rechts-click um Inhaltsverzeichnis hinzuzufügen - Update Feld)"

        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)
        p_element = paragraph._p
        document.add_page_break()

    # Writing paragraphs
    for index, row in data.iterrows():
        document.add_heading(f"{row['PARAGRAPH']} - {row['PARAGRAPH_TITLE']}", level=len(row['PARAGRAPH'].replace('.', '')))
        paragraph = document.add_paragraph()
        paragraph.add_run(f"\n{row['PARAGRAPH_TEXT']}\n")
        paragraph.style.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        document.add_page_break()

    # Adding summary
    counter = int(last_chapter)
    if paragraph_of_summary:
        counter += 1
        document.add_heading(f"{counter} - Zusammenfassung", level=1)
        paragraph = document.add_paragraph('Zusammenfassung')
        document.add_page_break()

    # Adding glossar
    if table_of_glossar:
        counter += 1
        document.add_heading(f"{counter} - Glossar", level=1)
        paragraph = document.add_paragraph('Glossar')
        document.add_page_break()

    # Adding stakeholders
    if table_of_stakeholders:
        counter += 1
        document.add_heading(f"{counter} - Stakeholder", level=1)
        paragraph = document.add_paragraph('Stakeholder')
        document.add_page_break()

    # Adding attachments
    if table_of_attachments:
        counter += 1
        document.add_heading(f"{counter} - Anhänge", level=1)
        paragraph = document.add_paragraph('Anhänge')

    # Download button
    buffer = BytesIO()
    document.save(buffer)
    st.toast('Das Dokument ist fertig!', icon ='📃')
    st.download_button(label='Download Template', data=buffer, file_name='BAS_Anzeige_Template.docx', mime='application/vnd.openxmlformats', icon='📃')
